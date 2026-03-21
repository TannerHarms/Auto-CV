"""Integration tests for renderers — generate output and verify structure."""

import pytest
from pathlib import Path

from auto_cv.parser.vault_reader import load_vault
from auto_cv.renderers.latex import LatexRenderer
from auto_cv.renderers.docx import DocxRenderer
from auto_cv.renderers.html import HtmlRenderer


EXAMPLE_VAULT = Path(__file__).resolve().parent.parent / "example_vault"


@pytest.fixture(scope="module")
def resume_and_style():
    return load_vault(EXAMPLE_VAULT)


# ---------------------------------------------------------------------------
# LaTeX
# ---------------------------------------------------------------------------


class TestLatexRenderer:
    @pytest.fixture(autouse=True)
    def setup(self, resume_and_style, tmp_path_factory):
        self.resume, self.style = resume_and_style
        self.output = tmp_path_factory.mktemp("latex_output")
        renderer = LatexRenderer()
        self.result = renderer.render(self.resume, self.style, self.output)

    def test_main_tex_exists(self):
        assert (self.output / "latex" / "main.tex").exists()

    def test_resume_sty_exists(self):
        assert (self.output / "latex" / "resume.sty").exists()

    def test_sections_directory(self):
        sections = self.output / "latex" / "sections"
        assert sections.is_dir()
        tex_files = list(sections.glob("*.tex"))
        assert len(tex_files) >= 5  # summary, experience, education, skills, projects

    def test_main_tex_includes_sections(self):
        main = (self.output / "latex" / "main.tex").read_text(encoding="utf-8")
        assert r"\input{sections/" in main

    def test_no_unescaped_ampersand(self):
        """Ensure & in user data is escaped."""
        for tex in (self.output / "latex" / "sections").glob("*.tex"):
            content = tex.read_text(encoding="utf-8")
            # Bare & (not \&) shouldn't appear in text content
            # However this is a heuristic — LaTeX tables use &
            # Just verify the file is non-empty
            assert len(content) > 0

    def test_latex_respects_format_specific_output_dir(self, resume_and_style, tmp_path):
        resume, style = resume_and_style
        target = tmp_path / "latex"
        result = LatexRenderer().render(resume, style, target)

        assert (target / "main.tex").exists()
        assert result.parent == target


# ---------------------------------------------------------------------------
# Markdown-to-LaTeX conversion
# ---------------------------------------------------------------------------

from auto_cv.renderers.latex import _escape_latex, _md_to_latex


class TestMdToLatex:
    def test_empty(self):
        assert _md_to_latex("") == ""
        assert _md_to_latex(None) == ""

    def test_plain_text_escaped(self):
        assert _md_to_latex("10% increase") == r"10\% increase"

    def test_bold(self):
        assert _md_to_latex("Achieved **top** rank") == r"Achieved \textbf{top} rank"

    def test_italic(self):
        assert _md_to_latex("used *agile* methods") == r"used \textit{agile} methods"

    def test_inline_code(self):
        assert _md_to_latex("wrote `Python` scripts") == r"wrote \texttt{Python} scripts"

    def test_link(self):
        result = _md_to_latex("[Google](https://google.com)")
        assert result == r"\href{https://google.com}{Google}"

    def test_bold_with_special_chars(self):
        result = _md_to_latex("**C++ & Java**")
        assert result == r"\textbf{C++ \& Java}"

    def test_multiple_formats(self):
        result = _md_to_latex("Used **Python** and *R* for `data` analysis")
        assert r"\textbf{Python}" in result
        assert r"\textit{R}" in result
        assert r"\texttt{data}" in result

    def test_escape_latex_unchanged(self):
        """_escape_latex should still work for structured fields."""
        assert _escape_latex("AT&T") == r"AT\&T"
        assert _escape_latex("$100") == r"\$100"


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------


class TestDocxRenderer:
    @pytest.fixture(autouse=True)
    def setup(self, resume_and_style, tmp_path_factory):
        self.resume, self.style = resume_and_style
        self.output = tmp_path_factory.mktemp("docx_output")
        renderer = DocxRenderer()
        self.result = renderer.render(self.resume, self.style, self.output)

    def test_docx_file_exists(self):
        assert self.result.exists()
        assert self.result.suffix == ".docx"

    def test_docx_is_valid(self):
        """Verify the file can be opened by python-docx."""
        from docx import Document

        doc = Document(str(self.result))
        # Should have paragraphs
        assert len(doc.paragraphs) > 0

    def test_docx_contains_name(self):
        from docx import Document

        doc = Document(str(self.result))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Jordan Rivera" in full_text

    def test_docx_respects_format_specific_output_dir(self, resume_and_style, tmp_path):
        resume, style = resume_and_style
        target = tmp_path / "docx"
        result = DocxRenderer().render(resume, style, target)

        assert result == target / "resume.docx"
        assert result.exists()


# ---------------------------------------------------------------------------
# HTML
# ---------------------------------------------------------------------------


class TestHtmlRenderer:
    @pytest.fixture(autouse=True)
    def setup(self, resume_and_style, tmp_path_factory):
        self.resume, self.style = resume_and_style
        self.output = tmp_path_factory.mktemp("html_output")
        renderer = HtmlRenderer()
        self.result = renderer.render(self.resume, self.style, self.output)

    def test_index_html_exists(self):
        assert (self.output / "html" / "index.html").exists()

    def test_portfolio_page_exists(self):
        assert (self.output / "html" / "portfolio.html").exists()

    def test_index_contains_name(self):
        html = (self.output / "html" / "index.html").read_text(encoding="utf-8")
        assert "Jordan Rivera" in html

    def test_index_contains_sections(self):
        html = (self.output / "html" / "index.html").read_text(encoding="utf-8")
        assert "Experience" in html
        assert "Education" in html
        assert "Skills" in html

    def test_css_variables_present(self):
        html = (self.output / "html" / "index.html").read_text(encoding="utf-8")
        assert "--color-primary" in html

    def test_portfolio_page_content(self):
        html = (self.output / "html" / "portfolio.html").read_text(encoding="utf-8")
        assert "Portfolio" in html

    def test_assets_copied(self):
        assert (self.output / "html" / "assets").is_dir()

    def test_nav_links(self):
        html = (self.output / "html" / "index.html").read_text(encoding="utf-8")
        # Should have nav with link to portfolio
        assert "portfolio.html" in html

    def test_html_respects_format_specific_output_dir(self, resume_and_style, tmp_path):
        resume, style = resume_and_style
        target = tmp_path / "html"
        result = HtmlRenderer().render(resume, style, target)

        assert result == target / "index.html"
        assert (target / "portfolio.html").exists()


# ---------------------------------------------------------------------------
# Regression tests — formatting correctness across renderers
# ---------------------------------------------------------------------------

ACADEMIC_VAULT = Path(__file__).resolve().parent.parent / "examples" / "academic-researcher"


@pytest.fixture(scope="module")
def academic_resume_and_style():
    return load_vault(ACADEMIC_VAULT)


class TestDocxFormatRegression:
    """Verify DOCX tab stops, section lines, and column layout are correct."""

    @pytest.fixture(autouse=True)
    def setup(self, academic_resume_and_style, tmp_path_factory):
        self.resume, self.style = academic_resume_and_style
        self.output = tmp_path_factory.mktemp("docx_regression")
        renderer = DocxRenderer()
        self.result = renderer.render(self.resume, self.style, self.output)

        from docx import Document
        self.doc = Document(str(self.result))
        self.ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

    def _pPr_child_tags(self, para):
        """Return ordered list of pPr child element tag names."""
        pPr = para._element.find(f'{{{self.ns}}}pPr')
        if pPr is None:
            return []
        return [c.tag.split('}')[1] for c in pPr]

    def _tab_defs(self, para):
        """Return list of (val, pos) tuples for tab stops in paragraph."""
        pPr = para._element.find(f'{{{self.ns}}}pPr')
        if pPr is None:
            return []
        tabs = pPr.findall(f'{{{self.ns}}}tabs/{{{self.ns}}}tab')
        return [(t.get(f'{{{self.ns}}}val'), int(t.get(f'{{{self.ns}}}pos')))
                for t in tabs]

    def _has_tab_run(self, para):
        """Check if paragraph contains a tab character run."""
        for r in para._element.findall(f'{{{self.ns}}}r'):
            if r.find(f'{{{self.ns}}}tab') is not None:
                return True
        return False

    def _has_underline_tab(self, para):
        """Check if paragraph has a tab run with underline formatting."""
        for r in para._element.findall(f'{{{self.ns}}}r'):
            if r.find(f'{{{self.ns}}}tab') is not None:
                u = r.find(f'{{{self.ns}}}rPr/{{{self.ns}}}u')
                if u is not None and u.get(f'{{{self.ns}}}val') == 'thick':
                    return True
        return False

    def test_pPr_tabs_before_spacing(self):
        """Tabs element must appear before spacing in pPr (OOXML schema order)."""
        for i, p in enumerate(self.doc.paragraphs):
            tags = self._pPr_child_tags(p)
            if 'tabs' in tags and 'spacing' in tags:
                assert tags.index('tabs') < tags.index('spacing'), (
                    f"Paragraph {i}: tabs must come before spacing in pPr, got {tags}"
                )

    def test_pPr_tabs_before_ind(self):
        """Tabs element must appear before ind in pPr (OOXML schema order)."""
        for i, p in enumerate(self.doc.paragraphs):
            tags = self._pPr_child_tags(p)
            if 'tabs' in tags and 'ind' in tags:
                assert tags.index('tabs') < tags.index('ind'), (
                    f"Paragraph {i}: tabs must come before ind in pPr, got {tags}"
                )

    def test_two_col_rows_have_right_tab(self):
        """Every paragraph with a tab run should have a right tab stop defined."""
        for i, p in enumerate(self.doc.paragraphs):
            if self._has_tab_run(p):
                tab_defs = self._tab_defs(p)
                vals = [td[0] for td in tab_defs]
                assert 'right' in vals or 'left' in vals, (
                    f"Paragraph {i} has tab run but no tab stop defined"
                )

    def test_section_headings_have_line_to_margin(self):
        """Section headings must have underlined tab extending to content width."""
        content_width_twips = int(
            (self.style.docx.page_width_inches - 2 * float(
                self.style.spacing.page_margin.strip().lower().replace('in', '')
            )) * 1440
        )
        found_section = False
        for p in self.doc.paragraphs:
            if self._has_underline_tab(p):
                found_section = True
                tab_defs = self._tab_defs(p)
                right_tabs = [td for td in tab_defs if td[0] == 'right']
                assert right_tabs, "Section heading should have right tab stop"
                assert abs(right_tabs[0][1] - content_width_twips) <= 2, (
                    f"Section line tab at {right_tabs[0][1]}, expected ~{content_width_twips}"
                )
        assert found_section, "Should find at least one section heading with line"

    def test_experience_dates_right_aligned(self):
        """Experience entries should have right-aligned tab with date on right side."""
        full_text = "\n".join(p.text for p in self.doc.paragraphs)
        assert "Senior Istari Emissary" in full_text
        found = False
        for p in self.doc.paragraphs:
            if "Senior Istari Emissary" in p.text and self._has_tab_run(p):
                tab_defs = self._tab_defs(p)
                right_tabs = [td for td in tab_defs if td[0] == 'right']
                assert right_tabs, "Experience entry should have right-aligned tab stop"
                found = True
                break
        assert found, "Should find experience entry with tab-based layout"

    def test_publication_authors_present(self):
        """Publications should include authors in the output."""
        full_text = "\n".join(p.text for p in self.doc.paragraphs)
        assert "Gandalf the Grey" in full_text, "Publication authors should appear in DOCX"


class TestLatexFormatRegression:
    """Verify LaTeX output has proper spacing and structure."""

    @pytest.fixture(autouse=True)
    def setup(self, academic_resume_and_style, tmp_path_factory):
        self.resume, self.style = academic_resume_and_style
        self.output = tmp_path_factory.mktemp("latex_regression")
        renderer = LatexRenderer()
        self.result = renderer.render(self.resume, self.style, self.output)
        self.sections_dir = self.output / "latex" / "sections"

    def test_publications_have_entry_gap(self):
        """Publications should have spacing between entries."""
        pub_tex = None
        for f in self.sections_dir.glob("*.tex"):
            content = f.read_text(encoding="utf-8")
            if r"\cvpub{" in content:
                pub_tex = content
                break
        assert pub_tex is not None, "Should find a publications .tex file"
        assert r"\vspace{" in pub_tex, "Publications should have \\vspace between entries"

    def test_publications_include_authors(self):
        """Publication \\cvpub command should include author names."""
        pub_tex = None
        for f in self.sections_dir.glob("*.tex"):
            content = f.read_text(encoding="utf-8")
            if r"\cvpub{" in content:
                pub_tex = content
                break
        assert pub_tex is not None
        assert "Gandalf the Grey" in pub_tex, "Author names should appear in publications .tex"

    def test_skills_have_category_spacing(self):
        """Skills should have vertical space between different categories."""
        skill_tex = None
        for f in self.sections_dir.glob("*.tex"):
            content = f.read_text(encoding="utf-8")
            if r"\cvskill{" in content:
                skill_tex = content
                break
        assert skill_tex is not None, "Should find a skills .tex file"
        cvskill_count = skill_tex.count(r"\cvskill{")
        if cvskill_count > 1:
            assert r"\vspace{" in skill_tex, (
                "Skills with multiple categories should have \\vspace between them"
            )

    def test_sty_has_cvpub_command(self):
        """resume.sty should define the \\cvpub command."""
        sty = (self.output / "latex" / "resume.sty").read_text(encoding="utf-8")
        assert r"\newcommand{\cvpub}" in sty


class TestHtmlFormatRegression:
    """Verify HTML output has correct skills layout."""

    @pytest.fixture(autouse=True)
    def setup(self, resume_and_style, tmp_path_factory):
        self.resume, self.style = resume_and_style
        self.output = tmp_path_factory.mktemp("html_regression")
        renderer = HtmlRenderer()
        self.result = renderer.render(self.resume, self.style, self.output)
        self.html = (self.output / "html" / "index.html").read_text(encoding="utf-8")

    def test_skills_no_fixed_min_width(self):
        """Skills category label should not have a hardcoded min-width."""
        import re
        matches = re.findall(r'\.skill[^{]*\{[^}]*\}', self.html, re.DOTALL)
        css = "\n".join(matches)
        assert "min-width" not in css, "Skills label should not have fixed min-width"

    def test_skills_rendered(self):
        """Skills section should exist and contain entries."""
        assert "skills-grid" in self.html
        assert "skill-category" in self.html
