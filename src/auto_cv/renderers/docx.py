"""DOCX renderer — generates a .docx resume using python-docx.

Produces preset-aware formatting that mirrors each LaTeX preset style:
matching section headings, name rendering, entry layouts, and
accent-colored decorations.
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.opc.constants import RELATIONSHIP_TYPE as RT

from auto_cv.models.resume import Resume, Section, SectionType
from auto_cv.models.style import StyleConfig
from auto_cv.renderers.base import BaseRenderer


def _hex_to_rgb(hex_color: str) -> RGBColor:
    h = hex_color.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _parse_pt(val: str) -> float:
    """Extract numeric pt value from a string like '10pt', '1.5mm', '0.5in'."""
    s = val.strip().lower()
    if s.endswith("mm"):
        return float(s[:-2]) * 2.835
    if s.endswith("cm"):
        return float(s[:-2]) * 28.35
    if s.endswith("in"):
        return float(s[:-2]) * 72
    return float(s.replace("pt", ""))


def _set_paragraph_spacing(para, before: int = 0, after: int = 0, line: int | None = None) -> None:
    """Set exact paragraph spacing in twips (1pt = 20 twips)."""
    pPr = para._element.get_or_add_pPr()
    # Remove any existing spacing element to avoid duplicates
    old = pPr.find(qn("w:spacing"))
    if old is not None:
        pPr.remove(old)
    spacing = pPr.makeelement(qn("w:spacing"), {})
    spacing.set(qn("w:before"), str(before))
    spacing.set(qn("w:after"), str(after))
    if line is not None:
        spacing.set(qn("w:line"), str(line))
        spacing.set(qn("w:lineRule"), "exact")
    # Insert after <w:tabs> to maintain correct OOXML pPr child order
    tabs_el = pPr.find(qn("w:tabs"))
    if tabs_el is not None:
        tabs_el.addnext(spacing)
    else:
        pPr.append(spacing)


def _add_tab_stop(para, position_emu: int, alignment: WD_TAB_ALIGNMENT,
                  *, leader: str | None = None) -> None:
    """Add a tab stop to a paragraph at the given EMU position."""
    pPr = para._element.get_or_add_pPr()
    tabs = pPr.find(qn("w:tabs"))
    if tabs is None:
        tabs = pPr.makeelement(qn("w:tabs"), {})
        # Insert before <w:spacing>/<w:ind> to maintain correct OOXML order
        ref = pPr.find(qn("w:spacing"))
        if ref is None:
            ref = pPr.find(qn("w:ind"))
        if ref is not None:
            ref.addprevious(tabs)
        else:
            pPr.append(tabs)
    attrs = {
        qn("w:val"): {
            WD_TAB_ALIGNMENT.RIGHT: "right",
            WD_TAB_ALIGNMENT.LEFT: "left",
            WD_TAB_ALIGNMENT.CENTER: "center",
        }.get(alignment, "left"),
        qn("w:pos"): str(int(position_emu / 635)),  # EMU → twips
    }
    if leader:
        attrs[qn("w:leader")] = leader
    tab = tabs.makeelement(qn("w:tab"), attrs)
    tabs.append(tab)


def _set_run_font(run, font_name: str) -> None:
    """Set all four w:rFonts facets so Word doesn't fall back to theme fonts."""
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.insert(0, rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rFonts.set(qn(attr), font_name)


def _add_run(para, text: str, *, size: float = 10, bold: bool = False,
             italic: bool = False, color: RGBColor | None = None,
             small_caps: bool = False, font_name: str | None = None) -> None:
    """Add a formatted run to a paragraph."""
    run = para.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    if small_caps:
        run.font.small_caps = True
    if font_name:
        _set_run_font(run, font_name)


# Markdown → DOCX runs
_MD_PATTERN = re.compile(
    r"\*\*(.+?)\*\*"       # **bold**
    r"|\*(.+?)\*"          # *italic*
    r"|`(.+?)`"            # `code`
    r"|\[(.+?)\]\((.+?)\)" # [text](url)
)


def _add_hyperlink(para, url: str, text: str, *, size: float = 10,
                   color: RGBColor | None = None,
                   font_name: str | None = None) -> None:
    """Insert a clickable hyperlink run into *para*."""
    part = para.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

    hyperlink = para._element.makeelement(qn("w:hyperlink"), {qn("r:id"): r_id})
    r = para._element.makeelement(qn("w:r"), {})
    rPr = para._element.makeelement(qn("w:rPr"), {})
    rStyle = para._element.makeelement(qn("w:rStyle"), {qn("w:val"): "Hyperlink"})
    rPr.append(rStyle)
    if font_name:
        rFonts = para._element.makeelement(qn("w:rFonts"), {
            qn("w:ascii"): font_name, qn("w:hAnsi"): font_name,
        })
        rPr.append(rFonts)
    sz_el = para._element.makeelement(qn("w:sz"), {qn("w:val"): str(int(size * 2))})
    rPr.append(sz_el)
    if color:
        c_el = para._element.makeelement(qn("w:color"), {
            qn("w:val"): f"{color[0]:02X}{color[1]:02X}{color[2]:02X}",
        })
        rPr.append(c_el)
    u_el = para._element.makeelement(qn("w:u"), {qn("w:val"): "none"})
    rPr.append(u_el)
    r.append(rPr)
    t = para._element.makeelement(qn("w:t"), {})
    t.text = text
    t.set(qn("xml:space"), "preserve")
    r.append(t)
    hyperlink.append(r)
    para._element.append(hyperlink)


def _add_md_runs(para, text: str, *, size: float = 10,
                 color: RGBColor | None = None,
                 font_name: str | None = None) -> None:
    """Parse simple markdown (bold, italic, code, links) into Word runs."""
    pos = 0
    for m in _MD_PATTERN.finditer(text):
        # Plain text before this match
        if m.start() > pos:
            _add_run(para, text[pos:m.start()], size=size, color=color, font_name=font_name)
        if m.group(1) is not None:          # **bold**
            _add_run(para, m.group(1), size=size, bold=True, color=color, font_name=font_name)
        elif m.group(2) is not None:        # *italic*
            _add_run(para, m.group(2), size=size, italic=True, color=color, font_name=font_name)
        elif m.group(3) is not None:        # `code`
            _add_run(para, m.group(3), size=size, color=color, font_name=font_name)
        elif m.group(4) is not None:        # [text](url)
            _add_hyperlink(para, m.group(5), m.group(4), size=size, color=color, font_name=font_name)
        pos = m.end()
    # Trailing plain text
    if pos < len(text):
        _add_run(para, text[pos:], size=size, color=color, font_name=font_name)


class DocxRenderer(BaseRenderer):
    def render(self, resume: Resume, style: StyleConfig, output_dir: Path) -> Path:
        docx_dir = self.prepare_output_dir(output_dir, "docx")
        docx_dir.mkdir(parents=True, exist_ok=True)
        out_path = docx_dir / "resume.docx"

        doc = Document()
        self._setup_page(doc, style)
        self._render_header(doc, resume, style)

        for section in resume.ordered_sections():
            self._render_section(doc, section, style)

        # Auto CV attribution — remove this block to hide
        self._render_attribution(doc, style)

        doc.save(str(out_path))
        return out_path

    # ------------------------------------------------------------------
    # Attribution
    # ------------------------------------------------------------------

    def _render_attribution(self, doc: Document, style: StyleConfig) -> None:
        """Add an attribution page. Delete this page in Word to remove."""
        # Page break
        bp = doc.add_paragraph()
        bp.runs[0].add_break(WD_BREAK.PAGE) if bp.runs else bp.add_run().add_break(WD_BREAK.PAGE)
        _set_paragraph_spacing(bp, before=0, after=0)

        # Vertical centering via spacer
        spacer = doc.add_paragraph()
        _set_paragraph_spacing(spacer, before=6000, after=0)

        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        font_name = style.fonts.body
        gray = RGBColor(0x99, 0x99, 0x99)
        _add_run(para, "Resume formatted and compiled by ", size=9, color=gray, font_name=font_name)
        _add_hyperlink(para, "https://github.com/TannerHarms/Auto-CV", "Auto CV",
                       size=9, color=gray, font_name=font_name)

    # ------------------------------------------------------------------
    # Page setup
    # ------------------------------------------------------------------

    def _setup_page(self, doc: Document, style: StyleConfig) -> None:
        section = doc.sections[0]
        section.page_width = Inches(style.docx.page_width_inches)
        section.page_height = Inches(style.docx.page_height_inches)
        margin = self._parse_margin(style.spacing.page_margin)
        section.top_margin = margin
        section.bottom_margin = margin
        section.left_margin = margin
        section.right_margin = margin

        # Page-wide content width for tab-stop calculations (in EMU)
        self._content_width = (
            section.page_width - section.left_margin - section.right_margin
        )

        doc_style = doc.styles["Normal"]
        doc_style.font.name = style.fonts.body
        doc_style.font.size = Pt(_parse_pt(style.fonts.size_base))
        doc_style.font.color.rgb = _hex_to_rgb(style.colors.text)
        # Set all four rFonts facets on the Normal style so Word doesn't
        # fall back to theme-linked fonts (Calibri, etc.)
        rPr = doc_style.element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = rPr.makeelement(qn("w:rFonts"), {})
            rPr.insert(0, rFonts)
        for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
            rFonts.set(qn(attr), style.fonts.body)
        # Remove theme font references that override explicit names
        for attr in ("w:asciiTheme", "w:hAnsiTheme", "w:eastAsiaTheme", "w:cstheme"):
            if rFonts.get(qn(attr)) is not None:
                del rFonts.attrib[qn(attr)]
        # Tight default paragraph spacing
        doc_style.paragraph_format.space_before = Pt(0)
        doc_style.paragraph_format.space_after = Pt(0)

    @staticmethod
    def _parse_margin(margin_str: str) -> Inches:
        s = margin_str.strip().lower()
        if s.endswith("cm"):
            return Inches(float(s[:-2]) / 2.54)
        if s.endswith("mm"):
            return Inches(float(s[:-2]) / 25.4)
        if s.endswith("pt"):
            return Inches(float(s[:-2]) / 72)
        return Inches(float(s.replace("in", "")))

    # ------------------------------------------------------------------
    # Two-column tab-stop helper (replaces table-based layout)
    # ------------------------------------------------------------------

    def _add_two_col_row(self, doc: Document, left_runs, right_runs,
                         *, right_width_emu: int | None = None):
        """Add a two-column line using a right-aligned tab stop.

        *left_runs* and *right_runs* are lists of dicts with keys
        accepted by _add_run (text, size, bold, italic, color, etc.).
        Returns the paragraph so callers can adjust spacing.
        """
        p = doc.add_paragraph()
        _set_paragraph_spacing(p, before=0, after=0)

        # Right-aligned tab stop at the content width
        _add_tab_stop(p, self._content_width, WD_TAB_ALIGNMENT.RIGHT)

        # Left side runs
        for run_spec in left_runs:
            _add_run(p, **run_spec)

        # Tab + right side runs (only if there's right content)
        if right_runs:
            p.add_run("\t")
            for run_spec in right_runs:
                _add_run(p, **run_spec)

        return p

    # ------------------------------------------------------------------
    # Header
    # ------------------------------------------------------------------

    def _render_header(self, doc: Document, resume: Resume, style: StyleConfig) -> None:
        preset = style.preset
        accent = _hex_to_rgb(style.colors.accent)
        primary = _hex_to_rgb(style.colors.primary)
        text_color = _hex_to_rgb(style.colors.text)
        secondary = _hex_to_rgb(style.colors.secondary)
        heading_color = _hex_to_rgb(style.colors.heading)

        name_parts = resume.config.name.rsplit(" ", 1)
        first = name_parts[0] if len(name_parts) > 1 else ""
        last = name_parts[-1]

        # --- Name ---
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_paragraph_spacing(p, before=0, after=40)

        if preset == "awesome-cv":
            # First name light, last name bold
            if first:
                _add_run(p, first + " ", size=32, color=text_color,
                         font_name=style.fonts.heading)
            _add_run(p, last, size=32, bold=True, color=text_color,
                     font_name=style.fonts.heading)
        elif preset == "executive":
            _add_run(p, resume.config.name, size=26, bold=True,
                     small_caps=True, color=primary,
                     font_name=style.fonts.heading)
        elif preset == "modern":
            _add_run(p, resume.config.name, size=28, bold=True,
                     color=primary, font_name=style.fonts.heading)
        elif preset == "creative":
            _add_run(p, resume.config.name, size=36, bold=True,
                     color=primary, font_name=style.fonts.heading)
        elif preset == "elegant":
            _add_run(p, resume.config.name, size=24, small_caps=True,
                     color=primary, font_name=style.fonts.heading)
        elif preset == "technical":
            _add_run(p, resume.config.name, size=24, bold=True,
                     color=primary, font_name=style.fonts.mono or style.fonts.heading)
        elif preset == "academic":
            _add_run(p, resume.config.name, size=20, bold=True,
                     color=text_color, font_name=style.fonts.heading)
        elif preset == "minimal":
            _add_run(p, resume.config.name, size=22, small_caps=True,
                     color=heading_color, font_name=style.fonts.heading)
        else:
            # classic / fallback
            _add_run(p, resume.config.name, size=24, bold=True,
                     color=primary, font_name=style.fonts.heading)

        # --- Title ---
        if resume.config.title:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_paragraph_spacing(p, before=0, after=40)
            if preset in ("awesome-cv",):
                _add_run(p, resume.config.title, size=7.6, small_caps=True,
                         color=accent, font_name=style.fonts.body)
            elif preset == "executive":
                _add_run(p, resume.config.title, size=12, italic=True,
                         color=accent, font_name=style.fonts.body)
            elif preset in ("modern", "creative"):
                _add_run(p, resume.config.title, size=12,
                         color=accent, font_name=style.fonts.heading)
            elif preset == "elegant":
                _add_run(p, resume.config.title, size=10, italic=True,
                         color=accent, font_name=style.fonts.body)
            elif preset == "technical":
                _add_run(p, resume.config.title, size=10,
                         color=accent, font_name=style.fonts.mono or style.fonts.heading)
            elif preset == "academic":
                _add_run(p, resume.config.title, size=10, italic=True,
                         color=secondary, font_name=style.fonts.body)
            else:
                _add_run(p, resume.config.title, size=12, italic=True,
                         color=secondary, font_name=style.fonts.body)

        # --- Contact ---
        contact = resume.config.contact
        contact_items: list[tuple[str, str | None]] = []  # (display, url|None)
        if contact.email:
            contact_items.append((contact.email, f"mailto:{contact.email}"))
        if contact.phone:
            contact_items.append((contact.phone, None))
        if contact.location:
            contact_items.append((contact.location, None))
        if contact.linkedin:
            contact_items.append(
                (f"linkedin.com/in/{contact.linkedin}",
                 f"https://linkedin.com/in/{contact.linkedin}"))
        if contact.github:
            contact_items.append(
                (f"github.com/{contact.github}",
                 f"https://github.com/{contact.github}"))
        if contact.website:
            url = contact.website if contact.website.startswith("http") else f"https://{contact.website}"
            contact_items.append((contact.website, url))

        if contact_items:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_paragraph_spacing(p, before=0, after=60)
            contact_color = secondary if preset in ("executive", "elegant", "academic", "minimal") else text_color
            contact_font = style.fonts.heading
            contact_size = 6.8
            if preset == "technical":
                contact_font = style.fonts.mono or style.fonts.heading
            elif preset in ("academic", "minimal"):
                contact_size = 7
            for i, (display, url) in enumerate(contact_items):
                if i > 0:
                    _add_run(p, " | ", size=contact_size, color=contact_color,
                             font_name=contact_font)
                if url:
                    _add_hyperlink(p, url, display, size=contact_size,
                                   color=contact_color, font_name=contact_font)
                else:
                    _add_run(p, display, size=contact_size, color=contact_color,
                             font_name=contact_font)

    # ------------------------------------------------------------------
    # Sections
    # ------------------------------------------------------------------

    def _render_section(self, doc: Document, section: Section, style: StyleConfig) -> None:
        preset = style.preset
        accent = _hex_to_rgb(style.colors.accent)
        primary = _hex_to_rgb(style.colors.primary)
        gray = _hex_to_rgb(style.colors.secondary)
        heading_color = _hex_to_rgb(style.colors.heading)
        heading_size = _parse_pt(style.fonts.size_heading)
        section_before = int(_parse_pt(style.spacing.section_gap) * 20)
        section_after = int(_parse_pt(style.spacing.header_to_content) * 20)

        title_text = section.title

        if preset == "creative":
            # Full-width colored banner with white uppercase text
            self._render_section_creative(doc, title_text, primary, heading_size,
                                          section_before, section_after, style)
        elif preset == "executive":
            # Double rule: thin above, heading, thick below
            self._render_section_executive(doc, title_text, accent, heading_color,
                                           heading_size, section_before, section_after, style)
        elif preset == "modern":
            # Left colored bar + uppercase heading
            self._render_section_modern(doc, title_text, primary, heading_color,
                                        heading_size, section_before, section_after, style)
        elif preset == "elegant":
            # Centered, thin lines above and below
            self._render_section_elegant(doc, title_text, accent, heading_color,
                                         heading_size, section_before, section_after, style)
        elif preset == "technical":
            # "> " prefix + dotted rule
            self._render_section_technical(doc, title_text, accent, heading_color,
                                           heading_size, section_before, section_after, style)
        elif preset in ("academic", "classic"):
            # Bold heading + full-width rule underneath
            self._render_section_classic(doc, title_text, primary, heading_color,
                                         heading_size, section_before, section_after, style,
                                         uppercase=(preset == "academic"))
        elif preset == "minimal":
            # Simple heading + thin hairline underneath
            self._render_section_minimal(doc, title_text, accent, heading_color,
                                         heading_size, section_before, section_after, style)
        else:
            # awesome-cv / fallback: accent text + gray rule to right margin
            self._render_section_awesome(doc, title_text, accent, gray,
                                         heading_size, section_before, section_after, style)

        handler = {
            SectionType.EXPERIENCE: self._render_experience,
            SectionType.EDUCATION: self._render_education,
            SectionType.SKILLS: self._render_skills,
            SectionType.PROJECTS: self._render_projects,
            SectionType.CERTIFICATIONS: self._render_certifications,
            SectionType.PUBLICATIONS: self._render_publications,
            SectionType.AWARDS: self._render_awards,
            SectionType.SUMMARY: self._render_summary,
            SectionType.VOLUNTEER: self._render_experience,
            SectionType.SERVICE: self._render_service,
            SectionType.LANGUAGES: self._render_languages,
            SectionType.INTERESTS: self._render_skills,
            SectionType.REFERENCES: self._render_references,
        }.get(section.section_type, self._render_custom)

        handler(doc, section, style)

    # -- Section heading variants -----------------------------------------------

    def _render_section_awesome(self, doc, title, accent, gray, size, before, after, style):
        """awesome-cv: accent text + gray rule from end of title to right margin."""
        p = doc.add_paragraph()
        _set_paragraph_spacing(p, before=before, after=after)
        _add_run(p, title, size=size, bold=True, color=accent,
                 font_name=style.fonts.body)
        _add_tab_stop(p, self._content_width, WD_TAB_ALIGNMENT.RIGHT)
        gray_hex = f"{gray[0]:02X}{gray[1]:02X}{gray[2]:02X}"
        run = p.add_run("\t")
        run.font.color.rgb = gray
        rPr = run._element.get_or_add_rPr()
        u_el = rPr.makeelement(qn("w:u"), {
            qn("w:val"): "thick", qn("w:color"): gray_hex})
        rPr.append(u_el)

    def _render_section_classic(self, doc, title, primary, heading_color, size,
                                before, after, style, *, uppercase=False):
        """classic/academic: bold heading + full-width rule underneath."""
        p = doc.add_paragraph()
        _set_paragraph_spacing(p, before=before, after=after)
        text = title.upper() if uppercase else title
        _add_run(p, text, size=size, bold=True, color=heading_color,
                 font_name=style.fonts.body)
        # Full-width rule below
        rule = doc.add_paragraph()
        _set_paragraph_spacing(rule, before=0, after=int(after * 0.5))
        pPr = rule._element.get_or_add_pPr()
        pBdr = pPr.makeelement(qn("w:pBdr"), {})
        primary_hex = f"{primary[0]:02X}{primary[1]:02X}{primary[2]:02X}"
        bottom = pBdr.makeelement(qn("w:bottom"), {
            qn("w:val"): "single", qn("w:sz"): "6",
            qn("w:space"): "1", qn("w:color"): primary_hex})
        pBdr.append(bottom)
        pPr.append(pBdr)

    def _render_section_executive(self, doc, title, accent, heading_color, size,
                                  before, after, style):
        """executive: thin rule above, small-caps heading, thick rule below."""
        accent_hex = f"{accent[0]:02X}{accent[1]:02X}{accent[2]:02X}"
        # Thin rule above
        rule_top = doc.add_paragraph()
        _set_paragraph_spacing(rule_top, before=before, after=0)
        pPr = rule_top._element.get_or_add_pPr()
        pBdr = pPr.makeelement(qn("w:pBdr"), {})
        bottom = pBdr.makeelement(qn("w:bottom"), {
            qn("w:val"): "single", qn("w:sz"): "4",
            qn("w:space"): "1", qn("w:color"): accent_hex})
        pBdr.append(bottom)
        pPr.append(pBdr)
        # Heading text
        p = doc.add_paragraph()
        _set_paragraph_spacing(p, before=40, after=40)
        _add_run(p, title, size=size, bold=True, small_caps=True,
                 color=heading_color, font_name=style.fonts.body)
        # Thick rule below
        rule_bot = doc.add_paragraph()
        _set_paragraph_spacing(rule_bot, before=0, after=after)
        pPr = rule_bot._element.get_or_add_pPr()
        pBdr = pPr.makeelement(qn("w:pBdr"), {})
        bottom = pBdr.makeelement(qn("w:bottom"), {
            qn("w:val"): "single", qn("w:sz"): "12",
            qn("w:space"): "1", qn("w:color"): accent_hex})
        pBdr.append(bottom)
        pPr.append(pBdr)

    def _render_section_modern(self, doc, title, primary, heading_color, size,
                               before, after, style):
        """modern: left colored bar + uppercase heading."""
        p = doc.add_paragraph()
        _set_paragraph_spacing(p, before=before, after=after)
        primary_hex = f"{primary[0]:02X}{primary[1]:02X}{primary[2]:02X}"
        # Left border as colored bar
        pPr = p._element.get_or_add_pPr()
        pBdr = pPr.makeelement(qn("w:pBdr"), {})
        left = pBdr.makeelement(qn("w:left"), {
            qn("w:val"): "single", qn("w:sz"): "18",
            qn("w:space"): "4", qn("w:color"): primary_hex})
        pBdr.append(left)
        pPr.append(pBdr)
        _add_run(p, title.upper(), size=size, bold=True,
                 color=heading_color, font_name=style.fonts.heading)

    def _render_section_creative(self, doc, title, primary, size, before, after, style):
        """creative: full-width colored banner with white uppercase text."""
        p = doc.add_paragraph()
        _set_paragraph_spacing(p, before=before, after=after)
        primary_hex = f"{primary[0]:02X}{primary[1]:02X}{primary[2]:02X}"
        # Shading (colored background)
        pPr = p._element.get_or_add_pPr()
        shd = pPr.makeelement(qn("w:shd"), {
            qn("w:val"): "clear", qn("w:color"): "auto",
            qn("w:fill"): primary_hex})
        pPr.append(shd)
        _add_run(p, title.upper(), size=size, bold=True,
                 color=RGBColor(0xFF, 0xFF, 0xFF), font_name=style.fonts.heading)

    def _render_section_elegant(self, doc, title, accent, heading_color, size,
                                before, after, style):
        """elegant: centered, thin decorative lines above and below."""
        accent_hex = f"{accent[0]:02X}{accent[1]:02X}{accent[2]:02X}"
        # Thin rule above
        rule_top = doc.add_paragraph()
        _set_paragraph_spacing(rule_top, before=before, after=0)
        pPr = rule_top._element.get_or_add_pPr()
        pBdr = pPr.makeelement(qn("w:pBdr"), {})
        bottom = pBdr.makeelement(qn("w:bottom"), {
            qn("w:val"): "single", qn("w:sz"): "2",
            qn("w:space"): "1", qn("w:color"): accent_hex})
        pBdr.append(bottom)
        pPr.append(pBdr)
        # Centered heading
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_paragraph_spacing(p, before=60, after=40)
        _add_run(p, title, size=size, bold=True, small_caps=True,
                 color=heading_color, font_name=style.fonts.body)
        # Thin rule below
        rule_bot = doc.add_paragraph()
        _set_paragraph_spacing(rule_bot, before=0, after=after)
        pPr = rule_bot._element.get_or_add_pPr()
        pBdr = pPr.makeelement(qn("w:pBdr"), {})
        bottom = pBdr.makeelement(qn("w:bottom"), {
            qn("w:val"): "single", qn("w:sz"): "2",
            qn("w:space"): "1", qn("w:color"): accent_hex})
        pBdr.append(bottom)
        pPr.append(pBdr)

    def _render_section_technical(self, doc, title, accent, heading_color, size,
                                  before, after, style):
        """technical: '> ' prefix + dotted rule to right margin."""
        p = doc.add_paragraph()
        _set_paragraph_spacing(p, before=before, after=after)
        mono = style.fonts.mono or style.fonts.heading
        _add_run(p, "> ", size=size, bold=True, color=accent, font_name=mono)
        _add_run(p, title.upper(), size=size, bold=True,
                 color=heading_color, font_name=mono)
        # Dotted rule after title
        _add_tab_stop(p, self._content_width, WD_TAB_ALIGNMENT.RIGHT, leader="dot")
        accent_hex = f"{accent[0]:02X}{accent[1]:02X}{accent[2]:02X}"
        run = p.add_run("\t")
        run.font.color.rgb = accent

    def _render_section_minimal(self, doc, title, accent, heading_color, size,
                                before, after, style):
        """minimal: simple small-caps heading + thin hairline."""
        p = doc.add_paragraph()
        _set_paragraph_spacing(p, before=before, after=0)
        _add_run(p, title.upper(), size=size, small_caps=True,
                 color=heading_color, font_name=style.fonts.heading)
        # Thin hairline
        rule = doc.add_paragraph()
        _set_paragraph_spacing(rule, before=0, after=after)
        pPr = rule._element.get_or_add_pPr()
        pBdr = pPr.makeelement(qn("w:pBdr"), {})
        accent_hex = f"{accent[0]:02X}{accent[1]:02X}{accent[2]:02X}"
        bottom = pBdr.makeelement(qn("w:bottom"), {
            qn("w:val"): "single", qn("w:sz"): "2",
            qn("w:space"): "1", qn("w:color"): accent_hex})
        pBdr.append(bottom)
        pPr.append(pBdr)

    # ------------------------------------------------------------------
    # Experience / Volunteer
    # ------------------------------------------------------------------

    def _render_experience(self, doc: Document, section: Section, style: StyleConfig) -> None:
        accent = _hex_to_rgb(style.colors.accent)
        dark = _hex_to_rgb(style.colors.text)
        gray = _hex_to_rgb(style.colors.secondary)
        entry_gap = int(_parse_pt(style.spacing.entry_gap) * 20)

        for i, entry in enumerate(section.experience_entries):
            # Row 1: title (left) | location (right, accent)
            left = [{"text": entry.title, "size": 10, "bold": True, "color": dark}]
            right = []
            if entry.location:
                right = [{"text": entry.location, "size": 9, "italic": True, "color": accent}]
            self._add_two_col_row(doc, left, right)

            # Row 2: organization (left, small-caps gray) | dates (right, gray)
            left = [{"text": entry.organization, "size": 8, "small_caps": True, "color": gray}]
            right = []
            if entry.dates:
                right = [{"text": entry.dates.display, "size": 8, "italic": True, "color": gray}]
            last_p = self._add_two_col_row(doc, left, right)

            # Bullet highlights
            bullet_size = _parse_pt(style.fonts.size_bullet)
            marker = "\u2013 " if style.spacing.bullet_marker == "dash" else "\u2022 "
            for h in entry.highlights:
                last_p = doc.add_paragraph()
                last_p.paragraph_format.left_indent = Inches(0.25)
                last_p.paragraph_format.first_line_indent = Inches(-0.15)
                _set_paragraph_spacing(last_p, before=0, after=0)
                _add_run(last_p, marker, size=bullet_size, color=dark)
                _add_md_runs(last_p, h, size=bullet_size, color=dark)

            # Entry gap: apply space-after on the last paragraph of this entry
            if i < len(section.experience_entries) - 1:
                _set_paragraph_spacing(last_p, before=0, after=entry_gap)

    # ------------------------------------------------------------------
    # Education
    # ------------------------------------------------------------------

    def _render_education(self, doc: Document, section: Section, style: StyleConfig) -> None:
        accent = _hex_to_rgb(style.colors.accent)
        dark = _hex_to_rgb(style.colors.text)
        gray = _hex_to_rgb(style.colors.secondary)
        entry_gap = int(_parse_pt(style.spacing.entry_gap) * 20)

        for i, entry in enumerate(section.education_entries):
            # Row 1: degree | location
            left = [{"text": entry.degree, "size": 10, "bold": True, "color": dark}]
            right = []
            if entry.location:
                right = [{"text": entry.location, "size": 9, "italic": True, "color": accent}]
            self._add_two_col_row(doc, left, right)

            # Row 2: institution | dates
            left = [{"text": entry.institution, "size": 8, "small_caps": True, "color": gray}]
            right = []
            if entry.dates:
                right = [{"text": entry.dates.display, "size": 8, "italic": True, "color": gray}]
            last_p = self._add_two_col_row(doc, left, right)

            # Description (plain text below header)
            bullet_size = _parse_pt(style.fonts.size_bullet)
            if entry.description:
                last_p = doc.add_paragraph()
                _set_paragraph_spacing(last_p, before=0, after=0)
                last_p.paragraph_format.left_indent = Inches(0.25)
                _add_md_runs(last_p, entry.description, size=bullet_size, color=dark)

            # GPA and highlights
            if entry.gpa:
                last_p = doc.add_paragraph()
                _set_paragraph_spacing(last_p, before=0, after=0)
                last_p.paragraph_format.left_indent = Inches(0.25)
                _add_run(last_p, f"GPA: {entry.gpa}", size=bullet_size, color=dark)

            marker = "\u2013 " if style.spacing.bullet_marker == "dash" else "\u2022 "
            for h in entry.highlights:
                last_p = doc.add_paragraph()
                last_p.paragraph_format.left_indent = Inches(0.25)
                last_p.paragraph_format.first_line_indent = Inches(-0.15)
                _set_paragraph_spacing(last_p, before=0, after=0)
                _add_run(last_p, marker, size=bullet_size, color=dark)
                _add_md_runs(last_p, h, size=bullet_size, color=dark)

            if i < len(section.education_entries) - 1:
                _set_paragraph_spacing(last_p, before=0, after=entry_gap)

    # ------------------------------------------------------------------
    # Skills / Interests
    # ------------------------------------------------------------------

    def _render_skills(self, doc: Document, section: Section, style: StyleConfig) -> None:
        dark = _hex_to_rgb(style.colors.text)
        text_color = _hex_to_rgb(style.colors.text)

        for cat in section.skill_categories:
            # Detect if skills are multi-line bold-label entries
            has_bold = any(s.strip().startswith("**") for s in cat.skills)
            if has_bold:
                # Render each skill entry as a separate line in the same row
                for j, skill in enumerate(cat.skills):
                    lbl = cat.name if j == 0 else ""
                    self._add_skill_row(doc, lbl, skill,
                                        label_color=dark, text_color=text_color)
            else:
                right_text = ", ".join(cat.skills)
                self._add_skill_row(doc, cat.name, right_text,
                                    label_color=dark, text_color=text_color)

    def _add_skill_row(self, doc: Document, label: str, skills_text: str,
                       *, label_color: RGBColor, text_color: RGBColor) -> None:
        """Skill row: right-aligned label, then left-aligned skills text with markdown."""
        label_w = Inches(1.5)
        gap = Inches(0.1)
        indent_twips = int((label_w + gap) / 635)  # EMU → twips

        p = doc.add_paragraph()
        _set_paragraph_spacing(p, before=0, after=0)

        # Indent wrapped lines to the second column
        pPr = p._element.get_or_add_pPr()
        ind = pPr.makeelement(qn("w:ind"), {
            qn("w:left"): str(indent_twips),
            qn("w:hanging"): str(indent_twips),
        })
        pPr.append(ind)

        # Right-aligned tab stop at label width for the label
        _add_tab_stop(p, label_w, WD_TAB_ALIGNMENT.RIGHT)
        # Left-aligned tab stop just after for the skills text
        _add_tab_stop(p, label_w + gap, WD_TAB_ALIGNMENT.LEFT)

        # Tab to the right-aligned stop, then the label
        p.add_run("\t")
        _add_run(p, label, size=10, bold=True, color=label_color)

        # Tab to the left-aligned stop, then skills with markdown
        p.add_run("\t")
        _add_md_runs(p, skills_text, size=9, color=text_color)

    # ------------------------------------------------------------------
    # Projects
    # ------------------------------------------------------------------

    def _render_projects(self, doc: Document, section: Section, style: StyleConfig) -> None:
        accent = _hex_to_rgb(style.colors.accent)
        dark = _hex_to_rgb(style.colors.text)
        gray = _hex_to_rgb(style.colors.secondary)
        entry_gap = int(_parse_pt(style.spacing.entry_gap) * 20)
        bullet_size = _parse_pt(style.fonts.size_bullet)

        for i, entry in enumerate(section.project_entries):
            # Row 1: name | dates (if any)
            left = [{"text": entry.name, "size": 10, "bold": True, "color": dark}]
            right = []
            if entry.dates:
                right = [{"text": entry.dates.display, "size": 9, "italic": True, "color": accent}]
            last_p = self._add_two_col_row(doc, left, right)

            # Technologies line
            if entry.technologies:
                last_p = doc.add_paragraph()
                _set_paragraph_spacing(last_p, before=0, after=0)
                _add_run(last_p, ", ".join(entry.technologies), size=8,
                         small_caps=True, color=gray)

            # Description
            if entry.description:
                last_p = doc.add_paragraph()
                _set_paragraph_spacing(last_p, before=0, after=0)
                _add_run(last_p, entry.description, size=bullet_size, color=dark)

            marker = "\u2013 " if style.spacing.bullet_marker == "dash" else "\u2022 "
            for h in entry.highlights:
                last_p = doc.add_paragraph()
                last_p.paragraph_format.left_indent = Inches(0.25)
                last_p.paragraph_format.first_line_indent = Inches(-0.15)
                _set_paragraph_spacing(last_p, before=0, after=0)
                _add_run(last_p, marker, size=bullet_size, color=dark)
                _add_md_runs(last_p, h, size=bullet_size, color=dark)

            if i < len(section.project_entries) - 1:
                _set_paragraph_spacing(last_p, before=0, after=entry_gap)

    # ------------------------------------------------------------------
    # Certifications
    # ------------------------------------------------------------------

    def _render_certifications(self, doc: Document, section: Section, style: StyleConfig) -> None:
        dark = _hex_to_rgb(style.colors.text)
        accent = _hex_to_rgb(style.colors.accent)

        for entry in section.certification_entries:
            # Single row: name + issuer (left) | date (right, accent)
            left = [{"text": entry.name, "size": 10, "bold": True, "color": dark}]
            if entry.issuer:
                left.append({"text": f" \u2013 {entry.issuer}", "size": 10, "color": dark})
            right = []
            if entry.date:
                right = [{"text": entry.date, "size": 9, "italic": True, "color": accent}]
            self._add_two_col_row(doc, left, right)

    # ------------------------------------------------------------------
    # Publications
    # ------------------------------------------------------------------

    def _render_publications(self, doc: Document, section: Section, style: StyleConfig) -> None:
        dark = _hex_to_rgb(style.colors.text)
        accent = _hex_to_rgb(style.colors.accent)
        gray = _hex_to_rgb(style.colors.secondary)
        entry_gap = int(_parse_pt(style.spacing.entry_gap) * 20)

        for i, entry in enumerate(section.publication_entries):
            # Row 1: title | date
            left = [{"text": entry.title, "size": 10, "bold": True, "color": dark}]
            right = []
            if entry.date:
                right = [{"text": entry.date, "size": 9, "italic": True, "color": accent}]
            last_p = self._add_two_col_row(doc, left, right)

            # Row 2: authors | venue
            left = []
            if entry.authors:
                left = [{"text": ", ".join(entry.authors), "size": 8,
                         "small_caps": True, "color": gray}]
            right = []
            if entry.venue:
                right = [{"text": entry.venue, "size": 8, "italic": True, "color": gray}]
            if left or right:
                last_p = self._add_two_col_row(doc, left, right)

            # Description
            if entry.description:
                last_p = doc.add_paragraph()
                _set_paragraph_spacing(last_p, before=0, after=0)
                last_p.paragraph_format.left_indent = Inches(0.25)
                _add_md_runs(last_p, entry.description, size=_parse_pt(style.fonts.size_bullet),
                             color=dark)

            if i < len(section.publication_entries) - 1:
                _set_paragraph_spacing(last_p, before=0, after=entry_gap)

    # ------------------------------------------------------------------
    # Awards
    # ------------------------------------------------------------------

    def _render_awards(self, doc: Document, section: Section, style: StyleConfig) -> None:
        dark = _hex_to_rgb(style.colors.text)
        accent = _hex_to_rgb(style.colors.accent)
        gray = _hex_to_rgb(style.colors.secondary)
        entry_gap = int(_parse_pt(style.spacing.entry_gap) * 20)
        bullet_size = _parse_pt(style.fonts.size_bullet)

        for i, entry in enumerate(section.award_entries):
            # Title + issuer (left) | date (right, accent)
            left = [{"text": entry.title, "size": 10, "bold": True, "color": dark}]
            if entry.issuer:
                left.append({"text": f", {entry.issuer}", "size": 9, "color": gray})
            right = []
            if entry.date:
                right = [{"text": entry.date, "size": 9, "italic": True, "color": accent}]
            last_p = self._add_two_col_row(doc, left, right)

            # Description
            if entry.description:
                last_p = doc.add_paragraph()
                _set_paragraph_spacing(last_p, before=0, after=0)
                _add_run(last_p, entry.description, size=bullet_size, color=dark)

            if i < len(section.award_entries) - 1:
                _set_paragraph_spacing(last_p, before=0, after=entry_gap)

    # ------------------------------------------------------------------
    # Service
    # ------------------------------------------------------------------

    def _render_service(self, doc: Document, section: Section, style: StyleConfig) -> None:
        dark = _hex_to_rgb(style.colors.text)
        accent = _hex_to_rgb(style.colors.accent)
        gray = _hex_to_rgb(style.colors.secondary)
        entry_gap = int(_parse_pt(style.spacing.entry_gap) * 20)
        bullet_size = _parse_pt(style.fonts.size_bullet)

        for i, entry in enumerate(section.experience_entries):
            # Title + org (left) | dates (right, accent)
            left = [{"text": entry.title, "size": 10, "bold": True, "color": dark}]
            if entry.organization:
                left.append({"text": f", {entry.organization}", "size": 8,
                             "small_caps": True, "color": gray})
            right = []
            if entry.dates:
                right = [{"text": entry.dates.display, "size": 9, "italic": True, "color": accent}]
            last_p = self._add_two_col_row(doc, left, right)

            # Highlights
            marker = "\u2013 " if style.spacing.bullet_marker == "dash" else "\u2022 "
            for h in entry.highlights:
                last_p = doc.add_paragraph()
                last_p.paragraph_format.left_indent = Inches(0.25)
                last_p.paragraph_format.first_line_indent = Inches(-0.15)
                _set_paragraph_spacing(last_p, before=0, after=0)
                _add_run(last_p, marker, size=bullet_size, color=dark)
                _add_md_runs(last_p, h, size=bullet_size, color=dark)

            if i < len(section.experience_entries) - 1:
                _set_paragraph_spacing(last_p, before=0, after=entry_gap)

    # ------------------------------------------------------------------
    # Summary
    # ------------------------------------------------------------------

    def _render_summary(self, doc: Document, section: Section, style: StyleConfig) -> None:
        if section.raw_content.strip():
            p = doc.add_paragraph()
            _set_paragraph_spacing(p, before=0, after=0)
            _add_run(p, section.raw_content.strip(), size=9,
                     color=_hex_to_rgb(style.colors.text))

    # ------------------------------------------------------------------
    # Languages
    # ------------------------------------------------------------------

    def _render_languages(self, doc: Document, section: Section, style: StyleConfig) -> None:
        dark = _hex_to_rgb(style.colors.text)
        text_color = _hex_to_rgb(style.colors.text)

        for entry in section.language_entries:
            self._add_skill_row(doc, entry.name, entry.proficiency or '',
                                label_color=dark, text_color=text_color)

    # ------------------------------------------------------------------
    # References
    # ------------------------------------------------------------------

    def _render_references(self, doc: Document, section: Section, style: StyleConfig) -> None:
        dark = _hex_to_rgb(style.colors.text)
        accent = _hex_to_rgb(style.colors.accent)
        gray = _hex_to_rgb(style.colors.secondary)
        entry_gap = int(_parse_pt(style.spacing.entry_gap) * 20)

        for i, entry in enumerate(section.reference_entries):
            # Name + relationship
            left = [{"text": entry.name, "size": 10, "bold": True, "color": dark}]
            if entry.relationship:
                left.append({"text": f" \u2013 {entry.relationship}", "size": 9,
                             "italic": True, "color": gray})
            last_p = self._add_two_col_row(doc, left, [])

            # Title at Organization
            parts: list[str] = []
            if entry.title:
                parts.append(entry.title)
            if entry.organization:
                parts.append(entry.organization)
            if parts:
                last_p = doc.add_paragraph()
                _set_paragraph_spacing(last_p, before=0, after=0)
                text = " at ".join(parts) if entry.title and entry.organization else parts[0]
                _add_run(last_p, text, size=8, small_caps=True, color=accent)

            # Contact info
            contact_parts: list[str] = []
            if entry.email:
                contact_parts.append(entry.email)
            if entry.phone:
                contact_parts.append(entry.phone)
            if contact_parts:
                last_p = doc.add_paragraph()
                _set_paragraph_spacing(last_p, before=0, after=0)
                _add_run(last_p, " | ".join(contact_parts), size=8, color=gray)

            if i < len(section.reference_entries) - 1:
                _set_paragraph_spacing(last_p, before=0, after=entry_gap)

    # ------------------------------------------------------------------
    # Fallback
    # ------------------------------------------------------------------

    def _render_custom(self, doc: Document, section: Section, style: StyleConfig) -> None:
        if section.raw_content.strip():
            for line in section.raw_content.strip().split("\n"):
                if line.strip():
                    p = doc.add_paragraph()
                    _set_paragraph_spacing(p, before=0, after=0)
                    _add_run(p, line, size=9, color=_hex_to_rgb(style.colors.text))
